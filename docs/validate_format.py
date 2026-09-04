"""
validate_format.py

Valida que un documento .docx cumpla con las reglas de formato del
Reglamento de Modalidades de Graduación de Institutos Técnicos y
Tecnológicos (RM 0487/2023, Artículo 9):

    - Fuente: Arial 11 en el cuerpo de texto
    - Alineación: justificado
    - Interlineado: 1.5 líneas
    - Márgenes: superior 3 cm, inferior 3 cm, izquierdo 3.5 cm, derecho 3 cm
    - Tamaño de papel: carta
    - Extensión: entre 30 y 80 páginas (aproximado por conteo de párrafos
      no vacíos, ya que python-docx no puede contar páginas reales sin
      abrir el documento en un motor de renderizado)

Uso:
    pip install python-docx --break-system-packages
    python validate_format.py mi_tesis.docx

El script no modifica el documento, solo reporta las desviaciones
encontradas.

Cómo extenderlo: a medida que confirmes las secciones marcadas con ⚠️ en
NORMAS_FORMATO.md (títulos, espaciado de párrafo, numeración de página,
norma de citas), agrega aquí una función `check_...()` nueva siguiendo el
mismo patrón que `check_body_paragraphs`.
"""

import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

# ---- Reglas básicas (Art. 9 del RM 0487/2023) --------------------------
EXPECTED_FONT_NAME = "Arial"
EXPECTED_FONT_SIZE_PT = 11
EXPECTED_ALIGNMENT = WD_ALIGN_PARAGRAPH.JUSTIFY
EXPECTED_LINE_SPACING = 1.5

# Márgenes en cm (Art. 9, inciso a)
EXPECTED_MARGIN_TOP_CM = 3.0
EXPECTED_MARGIN_BOTTOM_CM = 3.0
EXPECTED_MARGIN_LEFT_CM = 3.5
EXPECTED_MARGIN_RIGHT_CM = 3.0
MARGIN_TOLERANCE_CM = 0.05  # tolerancia por redondeo de unidades

MIN_PAGES = 30
MAX_PAGES = 80
# Aproximación: líneas de cuerpo por página con Arial 11 / interlineado 1.5
# en carta con estos márgenes. Es una estimación, no un conteo real de
# páginas — úsala solo como señal de alerta, no como verificación exacta.
APPROX_LINES_PER_PAGE = 27

# Nombres de estilo que se consideran "títulos" y por lo tanto se excluyen
# de la validación de cuerpo de texto. Ajusta si tu documento usa otros
# nombres de estilo (revisa con doc.styles).
HEADING_STYLE_PREFIXES = ("Heading", "Título", "Title", "Subtitle", "TOC")


def is_heading(paragraph) -> bool:
    style_name = (paragraph.style.name or "")
    return any(style_name.startswith(p) for p in HEADING_STYLE_PREFIXES)


def get_effective_font(paragraph):
    """
    Determina la fuente/tamaño 'efectivos' de un párrafo mirando sus runs.
    Si el párrafo no tiene runs con formato explícito, cae al estilo base.
    Devuelve (nombre_fuente, tamaño_pt) o (None, None) si no se puede
    determinar.
    """
    for run in paragraph.runs:
        if run.text.strip() == "":
            continue
        name = run.font.name
        size = run.font.size
        if name is not None or size is not None:
            size_pt = size.pt if size is not None else None
            return name, size_pt
    # Fallback: estilo del párrafo
    style_font = paragraph.style.font
    size_pt = style_font.size.pt if style_font.size else None
    return style_font.name, size_pt


def get_line_spacing(paragraph):
    pf = paragraph.paragraph_format
    return pf.line_spacing


def check_body_paragraphs(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "":
            continue
        if is_heading(para):
            continue

        snippet = (text[:50] + "…") if len(text) > 50 else text
        font_name, font_size = get_effective_font(para)
        alignment = para.alignment
        line_spacing = get_line_spacing(para)

        para_issues = []

        if font_name != EXPECTED_FONT_NAME:
            para_issues.append(
                f"fuente = {font_name!r} (se esperaba {EXPECTED_FONT_NAME!r})"
            )

        if font_size != EXPECTED_FONT_SIZE_PT:
            para_issues.append(
                f"tamaño = {font_size!r}pt (se esperaba {EXPECTED_FONT_SIZE_PT}pt)"
            )

        if alignment != EXPECTED_ALIGNMENT:
            para_issues.append(
                f"alineación = {alignment!r} (se esperaba JUSTIFY)"
            )

        if line_spacing != EXPECTED_LINE_SPACING:
            para_issues.append(
                f"interlineado = {line_spacing!r} (se esperaba {EXPECTED_LINE_SPACING})"
            )

        if para_issues:
            issues.append((i, snippet, para_issues))

    return issues


def check_margins(doc):
    """Revisa los márgenes de cada sección del documento contra el Art. 9."""
    issues = []
    for i, section in enumerate(doc.sections):
        checks = [
            ("superior", section.top_margin, EXPECTED_MARGIN_TOP_CM),
            ("inferior", section.bottom_margin, EXPECTED_MARGIN_BOTTOM_CM),
            ("izquierdo", section.left_margin, EXPECTED_MARGIN_LEFT_CM),
            ("derecho", section.right_margin, EXPECTED_MARGIN_RIGHT_CM),
        ]
        for name, value, expected_cm in checks:
            actual_cm = value.cm if value is not None else None
            if actual_cm is None or abs(actual_cm - expected_cm) > MARGIN_TOLERANCE_CM:
                issues.append(
                    f"Sección #{i}: margen {name} = {actual_cm!r} cm "
                    f"(se esperaba {expected_cm} cm)"
                )
    return issues


def check_page_size(doc):
    """Revisa que el tamaño de página sea carta (Letter: 21.59 x 27.94 cm)."""
    issues = []
    letter_width_cm = 21.59
    letter_height_cm = 27.94
    for i, section in enumerate(doc.sections):
        w = section.page_width.cm if section.page_width else None
        h = section.page_height.cm if section.page_height else None
        if w is None or h is None:
            continue
        if abs(w - letter_width_cm) > 0.2 or abs(h - letter_height_cm) > 0.2:
            issues.append(
                f"Sección #{i}: tamaño de página = {w:.2f}cm x {h:.2f}cm "
                f"(se esperaba carta: {letter_width_cm}cm x {letter_height_cm}cm)"
            )
    return issues


def estimate_page_count(doc):
    """
    Estimación aproximada de páginas contando párrafos no vacíos fuera de
    títulos. NO es un conteo real (eso requiere un motor de renderizado
    como Word o LibreOffice) — solo sirve como señal de alerta temprana.
    """
    non_empty = [p for p in doc.paragraphs if p.text.strip() and not is_heading(p)]
    return max(1, len(non_empty) // APPROX_LINES_PER_PAGE)


def main():
    if len(sys.argv) != 2:
        print("Uso: python validate_format.py archivo.docx")
        sys.exit(1)

    path = sys.argv[1]
    doc = Document(path)

    body_issues = check_body_paragraphs(doc)
    margin_issues = check_margins(doc)
    page_size_issues = check_page_size(doc)
    approx_pages = estimate_page_count(doc)

    any_issue = bool(body_issues) or bool(margin_issues) or bool(page_size_issues)

    if margin_issues:
        print(f"⚠️  Márgenes: {len(margin_issues)} desviación(es):")
        for m in margin_issues:
            print(f"    - {m}")
        print()

    if page_size_issues:
        print(f"⚠️  Tamaño de página: {len(page_size_issues)} desviación(es):")
        for p in page_size_issues:
            print(f"    - {p}")
        print()

    if body_issues:
        print(f"⚠️  Se encontraron {len(body_issues)} párrafo(s) con "
              f"desviaciones de fuente/alineación/interlineado:\n")
        for idx, snippet, para_issues in body_issues:
            print(f"Párrafo #{idx}: \"{snippet}\"")
            for problem in para_issues:
                print(f"    - {problem}")
            print()

    print(f"ℹ️  Extensión estimada: ~{approx_pages} página(s) "
          f"(estimación aproximada, no un conteo real). "
          f"El reglamento exige entre {MIN_PAGES} y {MAX_PAGES} páginas "
          f"(sin contar índice bibliográfico ni anexos).")
    if approx_pages < MIN_PAGES:
        print("    ⚠️  Podría estar por debajo del mínimo exigido — "
              "verificar con un conteo real en Word/LibreOffice.")
    elif approx_pages > MAX_PAGES:
        print("    ⚠️  Podría exceder el máximo exigido — "
              "verificar con un conteo real en Word/LibreOffice.")

    if not any_issue:
        print("\n✅ No se encontraron desviaciones en fuente, alineación, "
              "interlineado, márgenes ni tamaño de página.")
        return

    sys.exit(1)


if __name__ == "__main__":
    main()
